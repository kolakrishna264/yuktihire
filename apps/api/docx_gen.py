"""
DOCX Resume Generator
Uses python-docx to generate ATS-friendly Word documents.
"""
import asyncio
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_rule(doc):
    """Add a thin horizontal line below a paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_font(run, name="Arial", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, title: str):
    """Add a section heading with underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    set_font(run, size=11, bold=True)

    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '222222')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_docx_sync(resume_content: dict) -> bytes:
    """Synchronous DOCX generation."""
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    contact = resume_content.get("contact", {})
    name = contact.get("full_name") or resume_content.get("name", "")
    email = contact.get("email") or resume_content.get("email", "")
    phone = contact.get("phone") or resume_content.get("phone", "")
    location = contact.get("location") or resume_content.get("location", "")
    linkedin = contact.get("linkedin") or resume_content.get("linkedin", "")

    # Header: Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    set_font(name_run, size=20, bold=True)

    # Contact line
    contact_parts = [p for p in [email, phone, location, linkedin] if p]
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        set_font(contact_run, size=9.5, color=(68, 68, 68))

    doc.add_paragraph()  # spacer

    # Summary
    summary = resume_content.get("summary", "")
    if summary:
        add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        run = p.add_run(summary)
        set_font(run, size=10.5)

    # Experience
    experiences = resume_content.get("experiences", [])
    if experiences:
        add_section_heading(doc, "Professional Experience")
        for exp in experiences:
            # Company + Date row
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            company_run = p.add_run(exp.get("company", ""))
            set_font(company_run, bold=True, size=10.5)

            start = exp.get("start_date") or exp.get("startDate", "")
            end = exp.get("end_date") or exp.get("endDate", "")
            if exp.get("current"):
                end = "Present"
            dates = f"{start} – {end}" if start else ""

            if dates:
                tab_run = p.add_run("\t" + dates)
                set_font(tab_run, size=9.5, color=(80, 80, 80))
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))

            # Title
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(exp.get("title", ""))
            set_font(title_run, italic=True, size=10)

            # Bullets
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.2)
                run = bp.add_run(bullet.lstrip("•- "))
                set_font(run, size=10)

    # Skills
    skills = resume_content.get("skills", [])
    if skills:
        add_section_heading(doc, "Technical Skills")
        if isinstance(skills, str):
            p = doc.add_paragraph()
            run = p.add_run(skills)
            set_font(run, size=10.5)
        else:
            # Group by category
            by_category: dict[str, list] = {}
            for skill in skills:
                if isinstance(skill, dict):
                    cat = skill.get("category", "Skills")
                    by_category.setdefault(cat, []).append(skill.get("name", ""))
                else:
                    by_category.setdefault("Skills", []).append(str(skill))

            for cat, names in by_category.items():
                p = doc.add_paragraph()
                cat_run = p.add_run(f"{cat}: ")
                set_font(cat_run, bold=True, size=10)
                names_run = p.add_run(", ".join(names))
                set_font(names_run, size=10)

    # Education
    educations = resume_content.get("educations", [])
    if educations:
        add_section_heading(doc, "Education")
        for edu in educations:
            p = doc.add_paragraph()
            degree = f"{edu.get('degree', '')}"
            if edu.get("field"):
                degree += f", {edu['field']}"
            deg_run = p.add_run(degree)
            set_font(deg_run, bold=True, size=10.5)

            school_p = doc.add_paragraph()
            school_run = school_p.add_run(edu.get("school", ""))
            set_font(school_run, size=10)

            end_date = edu.get("end_date") or edu.get("endDate", "")
            gpa = edu.get("gpa", "")
            detail_parts = [p for p in [end_date, f"GPA: {gpa}" if gpa else ""] if p]
            if detail_parts:
                detail_p = doc.add_paragraph()
                detail_run = detail_p.add_run(" | ".join(detail_parts))
                set_font(detail_run, size=9.5, color=(80, 80, 80))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def generate_docx(resume_content: dict) -> bytes:
    """Async wrapper — runs sync generation in thread pool."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, generate_docx_sync, resume_content)
