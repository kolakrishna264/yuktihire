"""
Resume Parser — Extract text and structured data from PDF/DOCX uploads.
"""
import io
import re


async def parse_resume(content: bytes, filename: str) -> dict:
    """
    Extract structured data from resume file.
    Returns dict with name, email, phone, work experiences, educations, skills.
    """
    text = extract_text(content, filename)
    if not text:
        raise ValueError("Could not extract text from file")

    # Use AI to structure the extracted text
    from app.services.tailoring.jd_parser import extract_json_safe
    from anthropic import AsyncAnthropic
    from app.core.config import get_settings

    settings = get_settings()
    client = AsyncAnthropic(api_key=settings.anthropic_api_key)

    message = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=6000,
        system="You are a resume parser. Extract ALL data. Return ONLY valid JSON. Never truncate.",
        messages=[{
            "role": "user",
            "content": f"""Parse this resume and extract ALL information into JSON.

RESUME TEXT:
{text[:10000]}

CRITICAL: You MUST extract education completely. Look for:
- Degree names: B.S., B.Tech, M.S., M.Tech, MBA, PhD, Bachelor, Master
- University/college names
- Fields: Computer Science, Data Science, Engineering, etc.
- Graduation years
Education is often at the bottom of the resume. Extract ALL education entries with full details. Never return "Not specified" — look carefully at the resume text.

Return ONLY this JSON (education MUST be included):
{{
  "name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "github": "",
  "portfolio": "",
  "headline": "",
  "summary": "",
  "educations": [
    {{
      "degree": "Master of Science / Bachelor of Technology / etc",
      "field": "Computer Science / Data Science / etc",
      "school": "University name",
      "start_date": "YYYY or null",
      "end_date": "YYYY or null",
      "gpa": "3.8 or null"
    }}
  ],
  "certifications": [],
  "experiences": [
    {{
      "title": "job title",
      "company": "company name",
      "location": "",
      "start_date": "Mon YYYY",
      "end_date": "Mon YYYY or Present",
      "current": false,
      "bullets": ["bullet1", "bullet2"],
      "skills_used": []
    }}
  ],
  "skills": ["skill1", "skill2", "etc"],
  "projects": []
}}

IMPORTANT: Education section is REQUIRED. Extract ALL degrees. Do not skip education even if it appears at the bottom of the resume. For experience bullets, include max 4 per job to save space."""
        }]
    )

    raw_response = message.content[0].text
    print(f"[ResumeParser] AI response length: {len(raw_response)}")

    try:
        parsed = extract_json_safe(raw_response)
    except Exception:
        # Fallback: try direct JSON parse
        import json
        try:
            # Find JSON in the response
            start = raw_response.find("{")
            end = raw_response.rfind("}") + 1
            if start >= 0 and end > start:
                parsed = json.loads(raw_response[start:end])
            else:
                parsed = {}
        except Exception:
            parsed = {}

    # Ensure all expected fields exist
    parsed.setdefault("name", "")
    parsed.setdefault("email", "")
    parsed.setdefault("phone", "")
    parsed.setdefault("experiences", [])
    parsed.setdefault("educations", [])
    parsed.setdefault("skills", [])
    parsed.setdefault("certifications", [])
    parsed.setdefault("projects", [])

    # Fix education entries with "Not specified" — try to extract from raw text
    for edu in parsed.get("educations", []):
        if edu.get("degree", "") in ["", "Not specified", "not specified"]:
            edu_text = _extract_education_from_text(text)
            if edu_text:
                parsed["educations"] = edu_text
                break

    # Log what was extracted
    print(f"[ResumeParser] Extracted: {parsed.get('name','?')} | {len(parsed.get('experiences',[]))} exp | {len(parsed.get('educations',[]))} edu | {len(parsed.get('skills',[]))} skills")

    parsed["confidence"] = 0.85
    parsed["raw_text"] = text
    return parsed


def _extract_education_from_text(text: str) -> list[dict]:
    """Fallback: extract education using strict patterns that won't match random text."""
    import re
    educations = []
    lines = text.split("\n")

    # Only match FULL degree names — not abbreviations like "as", "ms", "ba" which are common words
    degree_patterns = [
        r"(Master\s+of\s+\w+|Master['']?s\s+(?:of|in)\s+\w+|M\.S\.\s+in\s+\w+)",
        r"(Bachelor\s+of\s+\w+|Bachelor['']?s\s+(?:of|in)\s+\w+|B\.S\.\s+in\s+\w+|B\.Tech)",
        r"(Ph\.?D\.?\s+in\s+\w+|Doctor\s+of\s+\w+)",
        r"(MBA\b)",
    ]

    for i, line in enumerate(lines):
        line_clean = line.strip()
        if not line_clean or len(line_clean) < 10:
            continue
        # Skip lines that look like bullet points / job descriptions
        if line_clean.startswith(("•", "-", "–", "▪", "*", "Built", "Led", "Designed", "Developed", "Created", "Implemented")):
            continue

        for pattern in degree_patterns:
            match = re.search(pattern, line_clean, re.IGNORECASE)
            if match:
                degree = match.group(0).strip()

                # Find field of study
                field = ""
                field_match = re.search(r"(?:in|of)\s+(.+?)(?:\s*[,|•\-–—]|\s+at\s+|\s+from\s+|\s*$)", line_clean[match.end():], re.IGNORECASE)
                if field_match:
                    field = field_match.group(1).strip()

                # Find school
                school = ""
                uni_match = re.search(r"(University|College|Institute|School|Academy)[\w\s,]+", line_clean, re.IGNORECASE)
                if uni_match:
                    school = uni_match.group(0).strip()
                elif i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if re.search(r"(University|College|Institute|School|Academy)", next_line, re.IGNORECASE):
                        school = next_line

                # Find year
                year_match = re.findall(r"20\d{2}", line_clean)
                end_date = year_match[-1] if year_match else None
                start_date = year_match[0] if len(year_match) > 1 else None

                educations.append({
                    "degree": degree,
                    "field": field or "",
                    "school": school or "",
                    "start_date": start_date,
                    "end_date": end_date,
                    "gpa": None,
                })
                break

    return educations if educations else []


def extract_text(content: bytes, filename: str) -> str:
    """Extract raw text from PDF or DOCX bytes."""
    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(content)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return extract_docx_text(content)
    else:
        # Try PDF first, then DOCX
        try:
            text = extract_pdf_text(content)
            if text.strip():
                return text
        except Exception:
            pass
        try:
            return extract_docx_text(content)
        except Exception:
            pass
    return ""


def extract_pdf_text(content: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_docx_text(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)
